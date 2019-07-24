import docx

guests_file = open('./data/guests.txt')
guests = guests_file.readlines()

blank = docx.Document('./data/template.docx')

for guest_index in range(len(guests)):
  if guest_index == 0:
    blank.paragraphs[0].text = 'It would be a pleasure to have the company of'
    blank.paragraphs[0].style = 'Style1'
  else:
    blank.add_paragraph(
        'It would be a pleasure to have the company of', 'Style1')

  blank.add_paragraph(guests[guest_index].rstrip(), 'Style2')
  blank.add_paragraph(
      'at 11010 Memory Lane on the Evening of', 'Style1')
  blank.add_paragraph('April 1st', 'Style3')
  blank.add_paragraph('at 7 o\'clock', 'Style1')

  if guest_index < len(guests) - 1:
    blank.add_paragraph().add_run().add_break(docx.text.run.WD_BREAK.PAGE)

blank.save('./dist/invitations.docx')
guests_file.close()
