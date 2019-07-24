import docx


def get_text(filename):
  doc = docx.Document(filename)

  full_text = []

  for paragraph in doc.paragraphs:
    full_text.append(paragraph.text)

  return '\n'.join(full_text)


doc = docx.Document('./data/demo.docx')

# Changing styles
doc.paragraphs[0].style = 'Quote'
doc.paragraphs[1].runs[0].style = 'Quote Char'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[2].strike = True

paragraph = doc.add_paragraph('Hello, World!')
paragraph.add_run(' ... and something that I forgot')

# Adding break
paragraph.runs[1].add_break(docx.text.run.WD_BREAK.PAGE)

doc.add_paragraph('Bye, World!')


# Adding heads
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)

doc.save('./dist/demo.docx')
